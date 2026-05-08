"""Annotate the customer SRS with Phase 1 gap comments.

Reads the May-2026 SRS, scans for Phase 1 commitments (July 2026 OPD go-live),
attaches Word comments where MedBrains has gaps, partials, or implementation
notes the management team needs to be aware of before signing off scope.

Output: ``presentations/SRS-Phase1-Gaps-Annotated.docx``.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document

SRC = Path(
    "/Users/apple/Downloads/Draft Software Requirements Specification 04302026 (1).docx"
)
OUT = Path(
    "/Users/apple/Projects/MedBrains/presentations/SRS-Phase1-Gaps-Annotated.docx"
)

AUTHOR = "MedBrains Build Team"
INITIALS = "MB"

# (regex_pattern, severity, comment_text). First match per paragraph wins.
RULES: list[tuple[str, str, str]] = [
    # ---- Phase 1 hard scope (Appendix H) ----
    (
        r"\bphysio\s*therapy\b|^\s*Physiotherapy\b",
        "GAP",
        "Phase 1 GAP — Physiotherapy module is not built. "
        "Required by Appendix H Phase 1 scope (OPD go-live July 2026). "
        "Needs: appointment scheduling, modality catalogue (TENS/IFT/US/exercise), "
        "session notes, treatment-plan templates, per-session and package billing, "
        "outcome tracking. Effort estimate: ~1 sprint (5–7 working days).",
    ),
    (
        r"\bDay\s*Care\s+Visit\s+Number\b|Register as Day Care Patient",
        "GAP",
        "Phase 1/2 GAP — Day-Care registration path (DCVN distinct from OPD/IPD) is not modelled. "
        "Currently only OPD visit + IPD admission flows exist. "
        "Needs new visit type, DCVN sequence, day-care package billing, "
        "discharge summary template. Effort: ~5 days.",
    ),
    (
        r"\bKiosk\s*&\s*Self-?Service|Self-?Service.*kiosk|kiosk.*self-?service",
        "GAP",
        "Phase 1 GAP — Self-service kiosk UI is not built. "
        "Required for token issue, registration check-in, walk-in queue join, "
        "kiosk-mode payment. Needs Kiosk shell (web or RN), restricted browser, "
        "thermal printer + Aadhaar reader integration. Effort: ~1 sprint.",
    ),
    (
        r"Walk-in registration with Aadhaar|Aadhaar/ID scanning",
        "PARTIAL",
        "Phase 1 PARTIAL — Aadhaar / ID scanner UI fields are ready, but device SDK "
        "(Mantra / Morpho / Cogent) is not embedded. ABDM Health-ID linking is live. "
        "Hardware integration effort: ~3 days per device family.",
    ),
    (
        r"Biometric/geo-fencing|biometric.*attendance|fingerprint",
        "PARTIAL",
        "PARTIAL — Biometric attendance UI hooks exist in HR module; device SDK "
        "(eSSL / Realtime / Mantra) and geo-fencing daemon are not deployed. "
        "Effort: ~5 days for device + 2 days for geo-fence policy.",
    ),
    (
        r"^\s*DICOM viewer access|9\.2 RIS & PACS|\bDICOM viewer\b",
        "PARTIAL",
        "Phase 1 PARTIAL — Radiology module has DICOM and HL7v2 message handlers, "
        "but PACS server (Orthanc / dcm4chee) integration and zero-footprint DICOM "
        "viewer (Cornerstone.js / OHIF) are not wired into the UI. "
        "Effort: ~1 sprint to embed viewer + storage path.",
    ),
    (
        r"SMS/WhatsApp(/Email)?\s+reminders|SMS reminders for follow-up",
        "PARTIAL",
        "Phase 1 PARTIAL — DLT-template registry, consent and content layer are built. "
        "Provider gateway wiring (Gupshup / Twilio / MSG91 / Karix) is partial — "
        "needs production credentials, retry queue, delivery receipts. "
        "Effort: ~2–3 days per gateway.",
    ),
    (
        r"Insurance pre-?auth|Pre-?Authoriz|preauth",
        "PARTIAL",
        "Phase 1 PARTIAL — ABDM/NHCX claim submission and TPA reconciliation are live. "
        "Pre-authorisation round-trip (request → query → authorisation letter) "
        "is data-modelled but not exposed in UI. Effort: ~1 sprint.",
    ),
    (
        r"Digital signature support for doctors|Digital signature by doctor|"
        r"Consent capture \(digital signature\)|Work order closure with digital signature",
        "PARTIAL",
        "PARTIAL — Signed-documents and signature-metadata tables are built. "
        "Signature pad / capacitive tablet driver and PKI signing certificate "
        "issuance flow are not deployed. Effort: ~3 days for pad + 5 days for PKI.",
    ),
    (
        r"Barcode specimen tracking|Mandatory barcode scanning|"
        r"Barcode-based medicine issue|Barcode-based stock movement|"
        r"Tray barcode tracking|Barcode labels|Barcode/QR-based counting",
        "PARTIAL",
        "PARTIAL — Barcode generation (label print) and table-side schema are ready. "
        "Handheld scanner driver (Zebra / Honeywell HID-keyboard mode) is supported, "
        "but ward-floor hardware procurement and per-station mapping are pending. "
        "Effort: small dev (~2 days), main effort is hardware rollout.",
    ),
    (
        r"\bBCMA\b|barcode medication administration|Bedside scanning|patient wristband",
        "PHASE-4",
        "Phase 4 (IPD) scope — BCMA and wristband scanning are scaffolded in nurse_mar. "
        "Hardware integration and workflow validation are required before IPD launch "
        "(Dec 2027 per Appendix H). Not a Phase 1 blocker.",
    ),
    (
        r"Queue status on mobile|Queue display|queue board",
        "PARTIAL",
        "Phase 1 PARTIAL — TV display app (apps/tv) and front-office queue stats "
        "endpoints are live. Per-OPD board configuration UI and WebSocket push "
        "from server are working in dev. Production hardware mapping pending. "
        "Effort: ~3–5 days.",
    ),
    (
        r"e-Prescription( auto-push)?|electronic prescription",
        "DONE",
        "DONE — CPOE order sets and e-prescription dispatch to pharmacy are built "
        "in OPD and IPD modules. Print + WhatsApp delivery hook pending pending "
        "the SMS/WhatsApp gateway item above.",
    ),
    # ---- Cross-cutting Phase 1 quality items ----
    (
        r"longitudinal record|Longitudinal records",
        "DONE",
        "DONE — patient_detail.tsx aggregates diagnoses, treatments, lab, imaging, "
        "medications, allergies and vitals across visits. Cross-encounter timeline "
        "view is implemented.",
    ),
    (
        r"HL7,?\s*FHIR,?\s*DICOM|FHIR R4",
        "DONE",
        "DONE — FHIR R4 routes (apps/web + crates/medbrains-server/src/routes/fhir.rs), "
        "ABDM/NHCX adapters, DICOM message handlers and HL7 v2 parsers are in place.",
    ),
    # ---- Phase 5 (Medical College) future scope ----
    (
        r"\bSkill Lab\b|Skill Lab & Simulation",
        "PHASE-5",
        "Phase 5 (Medical College) scope — Skill-lab inventory, scheduling and "
        "competency tracking are not built. NMC-mandated; planned for Q3–Q4 2026. "
        "Not a Phase 1 OPD blocker.",
    ),
    (
        r"Hostel\s*&\s*(Campus|transport)|Hostel\s*&\s*Campus Facilities|"
        r"Hostel\s*&\s*transport",
        "PHASE-5",
        "Phase 5 (Medical College) scope — Hostel allocation and transport-route "
        "management are not built. Required for NMC inspection. Effort: ~2 sprints. "
        "Not a Phase 1 OPD blocker.",
    ),
    (
        r"NMC compliance dashboard",
        "PHASE-5",
        "Phase 5 — Regulatory module captures the underlying data. NMC-specific "
        "dashboards (faculty workload, case-load exposure, attendance) need building "
        "before April 2027 medical-college submission.",
    ),
    (
        r"Day Surgery|DSU|Pre-op assessment",
        "PHASE-2",
        "Phase 2 (Sept 2026) scope — OT module + pre-op clearance forms exist; "
        "DSU-specific package billing, recovery monitoring and stand-alone DSU "
        "registration path need configuration. Effort: ~1 sprint.",
    ),
    # ---- Mobile App / patient-facing ----
    (
        r"Mobile App|Patient App|Doctor App",
        "PARTIAL",
        "PARTIAL — React Native CLI scaffold exists with shared API and types. "
        "Doctor-rounds, nurse-charting, patient bedside-portal and queue-status "
        "screens need feature build-out. Effort: ~2–3 sprints across roles.",
    ),
]


def find_runs_for_text(paragraph) -> list:
    """Return non-empty runs for anchoring a comment."""
    runs = [r for r in paragraph.runs if r.text]
    return runs


def add_summary_block(doc: Document) -> None:
    """Append a summary table to the front of the document."""
    # Insert before the first heading by adding to the body. python-docx
    # appends to end; we add at end to avoid layout disruption.
    doc.add_page_break()
    h = doc.add_heading("Phase 1 Readiness — MedBrains Self-Assessment", level=1)
    doc.add_paragraph(
        "Generated against the MedBrains in-house build, May 2026. "
        "Severity legend: GAP = not built; PARTIAL = partial / hardware pending; "
        "DONE = fully built; PHASE-2..5 = future scope per Appendix H."
    )

    summary = [
        ("Phase 1 — OPD Go-Live (July 2026)", "ON TRACK with stated gaps below."),
        ("Registration / OPD core",            "DONE"),
        ("Billing",                            "DONE (Phase 3 features)"),
        ("Pharmacy",                           "DONE (Phase 2 features, NDPS register live)"),
        ("Laboratory",                         "DONE (Phase 1–3 features)"),
        ("Radiology — orders & reporting",     "DONE"),
        ("Radiology — PACS viewer",            "PARTIAL — viewer embed pending"),
        ("Physiotherapy",                      "GAP — module not built"),
        ("Medical Camps",                      "DONE"),
        ("6 Clinical Departments",             "DONE — generic OPD covers; templates per-specialty optional"),
        ("Self-Service Kiosk",                 "GAP — kiosk shell not built"),
        ("SMS / WhatsApp gateway",             "PARTIAL — DLT done, provider wiring pending"),
        ("Aadhaar / Biometric devices",        "PARTIAL — UI ready, device SDK pending"),
        ("Digital signature pad",              "PARTIAL — PKI signing flow pending"),
        ("Mobile App (doctor / patient)",      "PARTIAL — scaffold ready, screens pending"),
        ("Insurance pre-authorisation",        "PARTIAL — claim flow live, preauth UI pending"),
        ("ABDM / NHCX",                        "DONE"),
        ("FHIR R4 / HL7 v2 / DICOM messaging", "DONE"),
        ("BCMA, wristbands",                   "PHASE 4 — IPD launch (Dec 2027)"),
        ("Skill lab / Hostel / NMC dashboards","PHASE 5 — Medical College readiness"),
    ]
    table = doc.add_table(rows=len(summary) + 1, cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "MedBrains Status"
    for r, (k, v) in enumerate(summary, 1):
        cells = table.rows[r].cells
        cells[0].text = k
        cells[1].text = v


def annotate(doc: Document) -> int:
    seen_per_rule: dict[int, int] = {}
    added = 0
    rules = [(re.compile(p, re.IGNORECASE), sev, txt) for p, sev, txt in RULES]
    for para in doc.paragraphs:
        text = para.text or ""
        if not text.strip():
            continue
        for idx, (rx, sev, body) in enumerate(rules):
            if rx.search(text):
                # Limit to 3 hits per rule so we do not spam the document.
                if seen_per_rule.get(idx, 0) >= 3:
                    break
                runs = find_runs_for_text(para)
                if not runs:
                    break
                comment_body = f"[{sev}] {body}"
                doc.add_comment(
                    runs=[runs[0], runs[-1]],
                    text=comment_body,
                    author=AUTHOR,
                    initials=INITIALS,
                )
                seen_per_rule[idx] = seen_per_rule.get(idx, 0) + 1
                added += 1
                break  # one comment per paragraph
    return added


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Source SRS not found: {SRC}")
    doc = Document(str(SRC))
    n = annotate(doc)
    add_summary_block(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Comments added: {n}")


if __name__ == "__main__":
    main()
